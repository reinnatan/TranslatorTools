from flask import Flask, render_template, request, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
from docx import Document
import os
import docx
import json

UPLOAD_FOLDER = './upload'
ALLOWED_ETENSIONS = {'docx', 'doc', 'pdf'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def hello():
	return 'Hello, World'
	
@app.route('/write-doc',methods=['POST'])
def write_doc():
	content = request.json
	to_loads = json.dumps(content, sort_keys=True, indent=3)
	resp = json.loads(to_loads)
	try:
		listToStr = ' '.join(map(str, resp['words'])) 
		write_doc(listToStr)
		return listToStr
	except KeyError as error:
		return jsonify({"message":"Key words tidak ditemukan"})


def write_doc(listToStr):
	print(listToStr)
	document = Document()
	document.add_paragraph(listToStr)
	document.save("./upload/testing.docx")


@app.route('/read-doc', methods=['POST'])
def read_doc():
	if request.method == "POST":
		try:
			f_name = UPLOAD_FOLDER+"/"+request.form['file']
			doc = docx.Document(f_name)
			all_paras = doc.paragraphs
			#print(docx.styles)
			fullText = []
			for para in all_paras:
				fullText.append(para.text)

			f =  open(f_name)
			return "Success read file : "+'\n'.join(fullText)
		except IOError as error:
			return error

#upload function
def allowed_file(filename):
	return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_ETENSIONS

@app.route('/uploads', methods=['GET','POST'])
def uploaad_file():
	if 'file' not in request.files:
		flash('No file part')
	file = request.files['file']
	print(file)

	if file and allowed_file(file.filename):
		filename = secure_filename(file.filename)
		file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
		#return 'Success filename'
		return redirect(url_for('success_upload_file',
                                    filename=filename))

@app.route('/success_uploaded_file', methods=['GET','POST'])
def success_upload_file():
	return "Success Uploaded file"

